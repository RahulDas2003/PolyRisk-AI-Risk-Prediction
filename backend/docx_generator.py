"""
DOCX Report Generator
Converts AI analysis JSON data to professional .docx reports
"""

import json
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

def add_color_to_cell(cell, color_hex):
    """Add background color to a table cell"""
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color_hex))
    cell._tc.get_or_add_tcPr().append(shading_elm)

def create_risk_color(risk_level):
    """Get color based on risk level"""
    colors = {
        'Low': '90EE90',      # Light green
        'Moderate': 'FFD700', # Gold
        'High': 'FF6B6B',     # Light red
        'Severe': 'DC143C'    # Crimson
    }
    return colors.get(risk_level, 'FFFFFF')

def generate_ai_analysis_docx(analysis_data, patient_name):
    """
    Generate a professional .docx report from AI analysis data
    
    Args:
        analysis_data: Dictionary containing AI analysis results
        patient_name: Name of the patient
    
    Returns:
        Document object ready to be saved
    """
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Title
    title = doc.add_heading('PolyRisk AI - Patient Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    subtitle = doc.add_paragraph(f'Patient: {patient_name}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.color.rgb = None
    
    # Report date
    date_para = doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.runs[0].font.size = Pt(10)
    date_para.runs[0].font.color.rgb = None
    
    # Add spacing
    doc.add_paragraph()
    
    # Extract data from analysis
    if 'analysis' in analysis_data:
        analysis = analysis_data['analysis']
    else:
        analysis = analysis_data
    
    # Risk Summary Section
    doc.add_heading('Risk Assessment Summary', level=1)
    
    if 'risk_summary' in analysis:
        risk_summary = analysis['risk_summary']
        
        # Overall Risk Score
        risk_table = doc.add_table(rows=1, cols=2)
        risk_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        risk_table.style = 'Table Grid'
        
        # Header row
        hdr_cells = risk_table.rows[0].cells
        hdr_cells[0].text = 'Overall Risk Score'
        hdr_cells[1].text = f"{risk_summary.get('overall_risk_score', 'N/A')}/10"
        
        # Style header
        for cell in hdr_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(12)
        
        # Risk level with color
        if 'risk_level' in risk_summary:
            risk_level = risk_summary['risk_level']
            color = create_risk_color(risk_level)
            add_color_to_cell(hdr_cells[1], color)
        
        # Risk notes
        if 'notes' in risk_summary:
            notes_para = doc.add_paragraph(f"Assessment Notes: {risk_summary['notes']}")
            notes_para.runs[0].font.size = Pt(11)
    
    # Base Risk Score and Scoring Breakdown
    if 'base_risk_score' in analysis:
        doc.add_heading('Risk Scoring Details', level=2)
        
        scoring_para = doc.add_paragraph(f"Base Risk Score: {analysis['base_risk_score']}/10")
        scoring_para.runs[0].font.bold = True
        
        if 'risk_summary' in analysis and 'scoring_breakdown' in analysis['risk_summary']:
            breakdown = analysis['risk_summary']['scoring_breakdown']
            
            breakdown_table = doc.add_table(rows=7, cols=2)
            breakdown_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            breakdown_table.style = 'Table Grid'
            
            breakdown_data = [
                ('Age Contribution', breakdown.get('age_contribution', 'N/A')),
                ('Kidney Function', breakdown.get('kidney_contribution', 'N/A')),
                ('Liver Function', breakdown.get('liver_contribution', 'N/A')),
                ('Drug Interactions', breakdown.get('drug_interactions', 'N/A')),
                ('Organ Effects', breakdown.get('organ_effects', 'N/A')),
                ('Polypharmacy', breakdown.get('polypharmacy', 'N/A'))
            ]
            
            for i, (label, value) in enumerate(breakdown_data):
                row = breakdown_table.rows[i]
                row.cells[0].text = label
                row.cells[1].text = str(value)
                
                # Style the label cell
                for paragraph in row.cells[0].paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True
    
    # Drug Analysis Section
    if 'drug_analysis' in analysis and analysis['drug_analysis']:
        doc.add_heading('Drug Analysis', level=1)
        
        for i, drug in enumerate(analysis['drug_analysis'], 1):
            # Drug name and category
            drug_heading = doc.add_heading(f'{i}. {drug.get("name", "Unknown Drug")}', level=2)
            
            if 'category' in drug:
                category_para = doc.add_paragraph(f"Category: {drug['category']}")
                category_para.runs[0].font.italic = True
                category_para.runs[0].font.size = Pt(11)
            
            # Individual risk score
            if 'individual_risk_score' in drug:
                risk_para = doc.add_paragraph(f"Individual Risk Score: {drug['individual_risk_score']}/10")
                risk_para.runs[0].font.bold = True
            
            # Drug interactions
            if 'interaction_risks' in drug and drug['interaction_risks']:
                doc.add_heading('Drug Interactions', level=3)
                
                for interaction in drug['interaction_risks']:
                    interaction_para = doc.add_paragraph(f"• With {interaction.get('drug', 'Unknown')}: {interaction.get('interaction', 'No details')}")
                    interaction_para.runs[0].font.size = Pt(10)
                    
                    if 'risk_score' in interaction:
                        risk_score = interaction['risk_score']
                        if risk_score >= 80:
                            interaction_para.runs[0].font.color.rgb = None  # Red
                        elif risk_score >= 60:
                            interaction_para.runs[0].font.color.rgb = None  # Orange
            
            # Side effects
            if 'side_effects' in drug and drug['side_effects']:
                doc.add_heading('Side Effects', level=3)
                
                for effect in drug['side_effects']:
                    if isinstance(effect, dict):
                        effect_text = f"• {effect.get('effect', 'Unknown')} ({effect.get('severity', 'Unknown')}, {effect.get('frequency', 'Unknown')})"
                    else:
                        effect_text = f"• {effect}"
                    
                    effect_para = doc.add_paragraph(effect_text)
                    effect_para.runs[0].font.size = Pt(10)
            
            # Organs affected
            if 'organs_affected' in drug and drug['organs_affected']:
                doc.add_heading('Organs Affected', level=3)
                
                for organ in drug['organs_affected']:
                    if isinstance(organ, dict):
                        organ_text = f"• {organ.get('organ', 'Unknown')} - {organ.get('effect', 'Unknown')} ({organ.get('severity', 'Unknown')})"
                    else:
                        organ_text = f"• {organ}"
                    
                    organ_para = doc.add_paragraph(organ_text)
                    organ_para.runs[0].font.size = Pt(10)
            
            # Add spacing between drugs
            doc.add_paragraph()
    
    # Drug Alternatives Section
    if 'drug_alternatives' in analysis and analysis['drug_alternatives']:
        doc.add_heading('Alternative Recommendations', level=1)
        
        for alt_group in analysis['drug_alternatives']:
            if 'original_drug' in alt_group:
                doc.add_heading(f"Alternatives for {alt_group['original_drug']}", level=2)
            
            if 'alternatives' in alt_group:
                for j, alternative in enumerate(alt_group['alternatives'], 1):
                    alt_heading = doc.add_heading(f"{j}. {alternative.get('alternative_name', 'Unknown Alternative')}", level=3)
                    
                    # Advantages
                    if 'advantages' in alternative and alternative['advantages']:
                        doc.add_heading('Advantages:', level=4)
                        for advantage in alternative['advantages']:
                            adv_para = doc.add_paragraph(f"• {advantage}")
                            adv_para.runs[0].font.size = Pt(10)
                    
                    # Disadvantages
                    if 'disadvantages' in alternative and alternative['disadvantages']:
                        doc.add_heading('Disadvantages:', level=4)
                        for disadvantage in alternative['disadvantages']:
                            dis_para = doc.add_paragraph(f"• {disadvantage}")
                            dis_para.runs[0].font.size = Pt(10)
                    
                    # Dosing recommendation
                    if 'dosing_recommendation' in alternative:
                        dosing_para = doc.add_paragraph(f"Dosing: {alternative['dosing_recommendation']}")
                        dosing_para.runs[0].font.bold = True
                        dosing_para.runs[0].font.size = Pt(10)
                    
                    # Monitoring parameters
                    if 'monitoring_parameters' in alternative and alternative['monitoring_parameters']:
                        doc.add_heading('Monitoring Required:', level=4)
                        for param in alternative['monitoring_parameters']:
                            param_para = doc.add_paragraph(f"• {param}")
                            param_para.runs[0].font.size = Pt(10)
                    
                    # Risk reduction
                    if 'risk_reduction' in alternative:
                        risk_reduction_para = doc.add_paragraph(f"Risk Reduction: {alternative['risk_reduction']}")
                        risk_reduction_para.runs[0].font.italic = True
                        risk_reduction_para.runs[0].font.size = Pt(10)
                    
                    doc.add_paragraph()  # Spacing between alternatives
    
    # Clinical Recommendations Section
    if 'clinical_recommendations' in analysis and analysis['clinical_recommendations']:
        doc.add_heading('Clinical Recommendations', level=1)
        
        for i, recommendation in enumerate(analysis['clinical_recommendations'], 1):
            rec_para = doc.add_paragraph(f"{i}. {recommendation}")
            rec_para.runs[0].font.size = Pt(11)
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph("This report was generated by PolyRisk AI - Advanced Drug Interaction Analysis System")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.runs[0].font.size = Pt(9)
    footer_para.runs[0].font.italic = True
    
    return doc

def save_analysis_to_docx(analysis_file_path, output_path):
    """
    Convert an AI analysis JSON file to DOCX format
    
    Args:
        analysis_file_path: Path to the JSON analysis file
        output_path: Path where the DOCX file should be saved
    
    Returns:
        bool: True if successful, False otherwise
    """
    try:
        # Load the analysis data
        with open(analysis_file_path, 'r', encoding='utf-8') as f:
            analysis_data = json.load(f)
        
        # Extract patient name from filename or data
        patient_name = "Unknown Patient"
        if 'patient_name' in analysis_data:
            patient_name = analysis_data['patient_name']
        elif 'analysis' in analysis_data and 'patient_name' in analysis_data['analysis']:
            patient_name = analysis_data['analysis']['patient_name']
        
        # Generate the document
        doc = generate_ai_analysis_docx(analysis_data, patient_name)
        
        # Save the document
        doc.save(output_path)
        
        return True
        
    except Exception as e:
        print(f"Error generating DOCX: {str(e)}")
        return False

if __name__ == "__main__":
    # Test the function
    test_file = "patient_analysis_data/ai_analysis_20251005_143002_Rahul_Das.json"
    output_file = "test_report.docx"
    
    if os.path.exists(test_file):
        success = save_analysis_to_docx(test_file, output_file)
        if success:
            print(f"SUCCESS: DOCX generated successfully: {output_file}")
        else:
            print("ERROR: Failed to generate DOCX")
    else:
        print(f"ERROR: Test file not found: {test_file}")
